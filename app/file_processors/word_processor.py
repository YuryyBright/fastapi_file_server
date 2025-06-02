import os
from app.file_processors.base_processor import FileProcessor, TextSearcher
from docx import Document
import zipfile
import olefile

class WordProcessor(FileProcessor):
    """
    Processor for Word files to perform text searches.
    Supports both .docx and .doc formats.
    """

    def __init__(self, file_path: str):
        """
        Initialize the Word processor with the file path.
        :param file_path: Path to the Word document.
        """
        self.file_path = file_path

    def search(self, keyword: str, exact_match: bool = True) -> list[str]:
        """
        Search for a keyword in the Word document text.
        :param keyword: The keyword to search for.
        :param exact_match: Whether to search for exact matches (True) or partial matches (False).
        :return: A list of lines containing the keyword.
        """
        if not os.path.exists(self.file_path):  # Check if file exists
            return [f"Error: File not found at {self.file_path}"]

        if self.is_docx(self.file_path):
            return self._search_docx(keyword, exact_match)
        elif self.is_doc(self.file_path):
            return self._search_doc(keyword, exact_match)
        else:
            return ["Error: Unsupported file type"]

    def is_docx(self, file_path: str) -> bool:
        """
        Check if the file is a .docx file.
        :param file_path: Path to the file.
        :return: True if the file is .docx, False otherwise.
        """
        return file_path.endswith('.docx')

    def is_doc(self, file_path: str) -> bool:
        """
        Check if the file is a .doc file.
        :param file_path: Path to the file.
        :return: True if the file is .doc, False otherwise.
        """
        return file_path.endswith('.doc')

    def _search_docx(self, keyword: str, exact_match: bool) -> list[str]:
        """
        Search within a .docx file using python-docx.
        :param keyword: The keyword to search for.
        :param exact_match: Whether to search for exact matches (True) or partial matches (False).
        :return: A list of lines containing the keyword.
        """
        try:
            doc = Document(self.file_path)
            results = []
            for para in doc.paragraphs:
                # Clean and append the paragraph text
                results.append(' '.join(para.text.split()))  # Remove multiple spaces
            return TextSearcher.search_text("\n".join(results), keyword, exact_match)
        except Exception as e:
            return [f"Error processing .docx file: {str(e)}"]

    def _search_doc(self, keyword: str, exact_match: bool) -> list[str]:
        """
        Search within a .doc file using the olefile library to extract text.
        :param keyword: The keyword to search for.
        :param exact_match: Whether to search for exact matches (True) or partial matches (False).
        :return: A list of lines containing the keyword.
        """
        try:
            ole = olefile.OleFileIO(self.file_path)
            if ole.exists('WordDocument'):
                doc_stream = ole.openstream('WordDocument')
                text = self._extract_text_from_doc(doc_stream)
                return TextSearcher.search_text(text, keyword, exact_match)
            else:
                return ["Error: Unable to extract text from .doc file"]
        except Exception as e:
            return [f"Error processing .doc file: {str(e)}"]

    def _extract_text_from_doc(self, doc_stream) -> str:
        """
        Extract text from a .doc file stream using olefile.
        :param doc_stream: The .doc file stream.
        :return: The extracted text from the .doc file.
        """
        # Extract raw text from the WordDocument stream (this is basic, you may need additional parsing)
        raw_data = doc_stream.read()
        return raw_data.decode('utf-16', errors='ignore')  # Basic decoding; refine if necessary

    def read(self) -> str:
        """
        Read the entire content of the Word document and return it as a single string.
        Supports both .docx and .doc formats.
        :return: The text content of the entire Word document.
        """
        if not os.path.exists(self.file_path):  # Check if file exists
            return f"Error: File not found at {self.file_path}"

        if self.is_docx(self.file_path):
            return self._read_docx()
        elif self.is_doc(self.file_path):
            return self._read_doc()
        else:
            return "Error: Unsupported file type"

    def _read_docx(self) -> str:
        """
        Read the entire content of a .docx file and return it as a single string.
        :return: The text content of the .docx file.
        """
        try:
            doc = Document(self.file_path)
            full_text = ""
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            return full_text.strip()
        except Exception as e:
            return f"Error processing .docx file: {str(e)}"

    def _read_doc(self) -> str:
        """
        Read the entire content of a .doc file and return it as a single string.
        :return: The text content of the .doc file.
        """
        try:
            ole = olefile.OleFileIO(self.file_path)
            if ole.exists('WordDocument'):
                doc_stream = ole.openstream('WordDocument')
                return self._extract_text_from_doc(doc_stream)
            else:
                return "Error: Unable to extract text from .doc file"
        except Exception as e:
            return f"Error processing .doc file: {str(e)}"
